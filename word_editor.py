import os
import shutil
from docx import Document


def substitute_File(old_word, new_words):
    new_words = new_words.split()
    files = os.listdir()
    files.remove("word_editor.py")
    for word in new_words:
        if not os.path.exists(word):
            os.makedirs(word)
        for file in files:
            try:
                shutil.copy(file, word + "/")
            except IOError as e:
                print("Unable to copy file. %s" % e)
                exit(1)
            document = Document(word+"/"+file)
            contains_old_word = 0
            for paragraph in document.paragraphs:
                if old_word in paragraph.text:
                    contains_old_word = 1
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        if old_word in inline[i].text:
                            text = inline[i].text.replace(old_word, word)
                            inline[i].text = text
                '''if contains_old_word == 0:
                    print("Can't find the text "+old_word+" in file " + file+" (源word里找不到你输入的关键词！)")'''
                document.save(word+"/"+file)

def searchWord(word):
    files = os.listdir()
    files.remove("word_editor.py")
    contains = 0
    for file in files:
        document = Document(file)
        for paragraph in document.paragraphs:
            if word in paragraph.text:
                print(" ")
                print("**********************************")
                print("Paragraph（段落）: "+ paragraph.text)
                print("File name（文件名）: "+ file)
                print("**********************************")
                contains = 1
    if contains == 0:
        print("No result found in files（当前根目录下没有文件包含目标关键词！）")

def main():
    print(" ")
    print(" ")
    print(" ")
    print("Welcome to word editor developed by Mingyong Cai. （欢迎来到Word Editor, 蔡明勇个人开发）")
    print("This is a tool to help you to process basic functions in Microsoft Word with multiple files. （这个程序将帮助你尝试多个Microsoft Word文件的关键词查找以及替代）")
    print("Which of the function you want to use?（请问要使用哪种功能？）")
    print(" ")
    print("1: Substitute multiple files（创建文件夹，替代根目录下word文档里的关键词）")
    print("2: Search keyword in multiple files（查找多个word文档里的关键词的位置）")
    inp = int(input("Enter a number（请输入数字）: "))

    if inp == 1:
        ori_word = str(input("Enter the old word you want to substitute（输入要替代的旧字符）: "))
        changed_word = str(input("Enter the new word, seperated by space（输入要替代的名字，空格隔开）: "))
        substitute_File(ori_word, changed_word)
    elif inp == 2:
        search_word = str(input("Enter the word you want to search（输入根目录下Word里你想查找的关键词）: "))
        searchWord(search_word)
    else:
        print("invalid input!")

    '''name = sys.argv[1]
    files = os.listdir()
    files.remove("substitute.py")
    for arg in sys.argv[2:]:
        if not os.path.exists(arg):
            os.makedirs(arg)
        for file in files:
            try:
                shutil.copy(file, arg+"/")
            except IOError as e:
                print("Unable to copy file. %s" % e)
                exit(1)'''



if __name__ == "__main__":
    main()
